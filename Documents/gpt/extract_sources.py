from pathlib import Path
from openpyxl import load_workbook
from docx import Document

base = Path(r"C:\GitLab\joboffer\Documents")
out = base / "gpt" / "source_extract.txt"
parts = []

wb = load_workbook(base / "Online Job Offer System - User Stories V2.xlsx", data_only=False)
for ws in wb.worksheets:
    parts.append(f"\n=== XLSX SHEET: {ws.title} ({ws.max_row}x{ws.max_column}) ===")
    for row in ws.iter_rows():
        vals = [str(c.value).strip() if c.value is not None else "" for c in row]
        if any(vals):
            parts.append(" | ".join(vals))

for name in ["Job Offer Business Logic Rules.docx", "Job Offer Programmer Detailed Flowchart - Editable Text.docx"]:
    doc = Document(base / name)
    parts.append(f"\n=== DOCX: {name} ===")
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for i, table in enumerate(doc.tables, 1):
        parts.append(f"-- TABLE {i} --")
        for row in table.rows:
            parts.append(" | ".join(c.text.strip().replace("\n", " / ") for c in row.cells))

out.write_text("\n".join(parts), encoding="utf-8")
print(out)
