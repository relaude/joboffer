from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = Path(r"C:\GitLab\joboffer\Documents")
OUT = BASE / "Online Job Offer System - Proposed Timeline.docx"

BLUE = "1F4E78"; MID = "5B9BD5"; LIGHT = "D9EAF7"; PALE = "EDF4F8"
DARK = "243746"; GRAY = "66737F"; WHITE = "FFFFFF"; GREEN = "70AD47"; GOLD = "FFC000"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for edge, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        el = tcMar.find(qn(f'w:{edge}'))
        if el is None: el=OxmlElement(f'w:{edge}'); tcMar.append(el)
        el.set(qn('w:w'), str(val)); el.set(qn('w:type'),'dxa')

def set_repeat(row):
    trPr=row._tr.get_or_add_trPr(); e=OxmlElement('w:tblHeader'); e.set(qn('w:val'),'true'); trPr.append(e)

def set_cell_text(cell, text, bold=False, color=DARK, size=8.5, align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align is not None: p.alignment=align
    p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0)
    r=p.add_run(str(text)); r.bold=bold; r.font.name='Aptos'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)

def table(doc, headers, rows, widths, font=8.5):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,(h,w) in enumerate(zip(headers,widths)):
        t.columns[i].width=Inches(w); t.rows[0].cells[i].width=Inches(w)
        set_cell_text(t.rows[0].cells[i],h,True,WHITE,font,WD_ALIGN_PARAGRAPH.CENTER); shade(t.rows[0].cells[i],BLUE)
    set_repeat(t.rows[0])
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,(v,w) in enumerate(zip(row,widths)):
            cells[i].width=Inches(w); set_cell_text(cells[i],v,False,DARK,font)
            if ri%2: shade(cells[i],PALE)
    return t

doc=Document(); sec=doc.sections[0]
sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
sec.header_distance=Inches(.3); sec.footer_distance=Inches(.3)
styles=doc.styles
normal=styles['Normal']; normal.font.name='Aptos'; normal.font.size=Pt(10); normal.font.color.rgb=RGBColor.from_string(DARK)
normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.08
for name,size,color,before,after in [('Title',26,BLUE,0,10),('Heading 1',16,BLUE,14,6),('Heading 2',12,MID,10,4)]:
    s=styles[name]; s.font.name='Aptos Display' if name!='Normal' else 'Aptos'; s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.font.bold=True
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

# Header/footer
h=sec.header.paragraphs[0]; h.text='ONLINE JOB OFFER SYSTEM  |  DELIVERY PLAN'; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in h.runs: r.font.name='Aptos'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=f.add_run('Confidential - Proposed timeline for planning purposes'); r.font.name='Aptos'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)

p=doc.add_paragraph(style='Title'); p.add_run('Proposed Implementation Timeline')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Online Job Offer System  |  One Full-Stack Developer  |  Baseline: 30 Weeks'); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor.from_string(MID)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Version 1.0  •  21 July 2026'); r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(GRAY)

doc.add_heading('Executive Summary',1)
doc.add_paragraph('A practical baseline is 30 calendar weeks from approved kickoff to production handover for one dedicated full-stack developer. The plan sequences the highest-risk foundations first—data model, access control, salary rules, proposal versioning, and approval routing—then completes candidate outcomes, documents, reporting, integration testing, UAT, and deployment.')

summary=[('Recommended baseline','30 weeks'),('Target MVP feature-complete','End of Week 23'),('System/UAT stabilization','Weeks 24-28'),('Production release','Week 30'),('Delivery capacity assumption','1 developer, ~80% planned delivery capacity'),('Estimate confidence','Planning estimate; refine after discovery and technical platform confirmation')]
table(doc,['Planning Item','Proposed Position'],summary,[2.2,4.8],9)

doc.add_heading('Scope Basis',1)
doc.add_paragraph('The estimate covers the end-to-end workflow described in the supplied sources: candidate intake through MS Forms, validation, legal-entity salary matrix selection, one-to-three proposal options, salary-band calculations, per-option approval decisions, President escalation for above-band offers, reminders/escalations, candidate discussion and re-proposal, offer-letter generation, delegation, audit history, governance reporting, and role-based access.')

doc.add_heading('Delivery Timeline at a Glance',1)
phases=[
('1. Discovery & architecture',1,2),('2. Platform foundation',3,5),('3. Candidate intake',6,8),('4. Salary analysis',9,12),('5. Approval workflow',13,17),('6. Candidate outcome & documents',18,20),('7. Reporting & administration',21,23),('8. System hardening',24,26),('9. UAT & release readiness',27,29),('10. Go-live & hypercare',30,30)]
headers=['Phase','W1-3','W4-6','W7-9','W10-12','W13-15','W16-18','W19-21','W22-24','W25-27','W28-30']
t=doc.add_table(rows=1,cols=11); t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.CENTER
widths=[2.05]+[.50]*10
for i,hv in enumerate(headers):
    t.columns[i].width=Inches(widths[i]); set_cell_text(t.rows[0].cells[i],hv,True,WHITE,7.2,WD_ALIGN_PARAGRAPH.CENTER); shade(t.rows[0].cells[i],BLUE)
for idx,(name,start,end) in enumerate(phases):
    cells=t.add_row().cells; set_cell_text(cells[0],name,False,DARK,7.6)
    for b in range(10):
        lo=b*3+1; hi=lo+2; active=not(end<lo or start>hi)
        set_cell_text(cells[b+1],'●' if active else '',True,WHITE,9,WD_ALIGN_PARAGRAPH.CENTER)
        shade(cells[b+1],MID if active else ('F5F7F9' if idx%2==0 else PALE))

doc.add_page_break()
doc.add_heading('Detailed Phase Plan',1)
rows=[
('1','W1-2','Discovery & solution architecture','Confirm platform, integrations, non-functional needs, workflow states, roles, data model, environments, backlog and acceptance approach.','Approved solution outline; prioritized backlog; architecture/data model; delivery setup','All; BL-01 to BL-20'),
('2','W3-5','Foundation, security & master data','Project skeleton, database migrations, authentication, RBAC/ABAC, legal entities, requisitions, users, approver matrix, salary-matrix administration and audit framework.','Working secured application; core master-data screens/APIs; CI/CD baseline','US-04, US-16, US-19; BL-01, 04-05, 17-18'),
('3','W6-8','Candidate intake & validation','Create job-offer case, MS Forms handoff/import approach, file validation/storage, submission tracking, correction loop and TA validation.','Traceable candidate intake; validated package workflow','US-01 to US-03'),
('4','W9-12','Salary analysis & proposal engine','Effective-dated matrix lookup; 1-3 proposal options; compa-ratio, increase and band-position calculations; exception flags; proposal version model.','Job Offer Analysis with reliable calculations and validation','US-04 to US-06, US-13, US-19; BL-05 to 09'),
('5','W13-17','Approval routing & notifications','Per-option decisions, HROD/Division approval, President escalation, rejection comments, secure task links, email/in-app notices, reminders, escalation, retries and delivery logs.','End-to-end approval workflow with SLA handling','US-07 to US-11, US-20, US-21; BL-10 to 12, 19'),
('6','W18-20','Candidate decision, re-proposal & offer letter','Approved-options-only discussion, accept/decline/negotiation, new proposal cycles with history, final offer merge and secure document storage.','Candidate outcome workflow; generated offer letter','US-12, US-13, US-15; BL-13 to 16'),
('7','W21-23','Delegation, dashboards & governance','Delegate/reassign/shadow access, transaction history, HRBP status view, above-band report, filters/export and administration refinements.','Feature-complete MVP; operational dashboards and reports','US-14, US-16 to US-18; BL-17, 18, 20'),
('8','W24-26','Integration, security & system testing','End-to-end scenarios, negative paths, authorization checks, audit reconciliation, file security, notification failures, performance checks, backups and defect repair.','Release candidate; system test evidence; operating notes','All stories and business rules'),
('9','W27-29','UAT, remediation & release readiness','Business-led UAT, developer fixes, regression, data/config preparation, user/admin guidance, deployment rehearsal and go-live approval.','UAT sign-off; deployment package; support runbook','All acceptance criteria'),
('10','W30','Production deployment & hypercare','Production release, smoke tests, monitoring, priority defect response, handover and backlog closure review.','Live system; handover; post-launch backlog','All')]
table(doc,['#','Weeks','Workstream','Key Activities','Exit Deliverable','Scope Trace'],rows,[.3,.55,1.15,2.15,1.55,1.3],7.6)

doc.add_page_break()
doc.add_heading('Milestones and Decision Gates',1)
milestones=[('M1','End W2','Scope and architecture approved','Platform/integration decisions and backlog are sufficiently clear to build.'),('M2','End W8','Intake workflow demo','TA can create a case, receive/validate candidate inputs, and see an auditable status history.'),('M3','End W12','Salary engine demo','Matrix lookup, calculations, band flags, and proposal versioning pass agreed examples.'),('M4','End W17','Approval workflow demo','Normal and above-band paths, per-option decisions, reminders, and escalation work end to end.'),('M5','End W23','Feature complete','All planned user stories are available in the test environment.'),('M6','End W26','Release candidate','System, security, and end-to-end tests pass with no open critical defects.'),('M7','End W29','Go-live approval','UAT signed off; configurations, deployment and support materials are ready.'),('M8','W30','Production handover','Smoke tests pass and ownership transfers to operations/support.')]
table(doc,['Gate','Target','Milestone','Exit Criteria'],milestones,[.55,.75,1.55,4.25],8.2)

doc.add_heading('Planning Assumptions',1)
assumptions=[
'The developer is allocated full time, with approximately 80% of capacity available for planned delivery after ceremonies, support and rework.',
'A product owner and representatives from TA, Total Rewards, HROD, HRBP, Division leadership and IT/security are available for weekly clarification and scheduled reviews.',
'A supported application platform, identity provider, email service, database, document storage and deployment environment are selected by the end of Week 2.',
'MS Forms integration is limited to a supported handoff/import/API pattern; building a new public candidate portal is outside this baseline.',
'One offer-letter template and one primary notification channel configuration are included initially; extensive template variants or Teams integration require re-estimation.',
'Business users supply approved salary matrices, approver mappings, SLA thresholds, legal entities, test data and offer templates on schedule.',
'UAT execution is performed by business users; the developer supports triage, remediation and regression.',
'Calendar duration excludes major organizational shutdowns and assumes stakeholder decisions are returned within two business days.'
]
for x in assumptions: doc.add_paragraph(x,style='List Bullet')

doc.add_heading('Dependencies and Schedule Risks',1)
risks=[
('High','Late platform/integration decisions','Delays architecture and can cause rework across intake, notifications and deployment.','Confirm technical stack and service owners in Week 1; time-box proof-of-concept work.'),
('High','Unavailable or changing salary/approval rules','Blocks calculation and workflow acceptance.','Obtain approved matrices and decision tables by Week 3; version all rule changes.'),
('High','Single-developer concentration','Illness, production support or competing work directly affects the critical path.','Protect allocation, document weekly, automate tests/deployment, and keep a visible contingency backlog.'),
('Medium','Identity, email, storage or MS Forms access','External approvals and tenant configuration can create idle time.','Raise access requests during discovery and use stubs/mocks until credentials arrive.'),
('Medium','Slow stakeholder review/UAT','Defects and rule gaps surface late.','Run demos at M2-M5 and book UAT participants before Week 23.'),
('Medium','Security/privacy findings','Compensation and candidate documents require strict access and auditability.','Threat-model early; test role boundaries and document access before UAT.')]
rt=table(doc,['Rating','Risk / Dependency','Schedule Effect','Mitigation'],risks,[.65,1.35,2.05,2.95],8)
for row in rt.rows[1:]:
    rating=row.cells[0].text; shade(row.cells[0], 'F4CCCC' if rating=='High' else 'FFF2CC')

doc.add_heading('Estimate Boundaries and Change Control',1)
doc.add_paragraph('The 30-week baseline is appropriate for planning and sequencing, but it is not a fixed-price commitment. Re-estimate after Week 2 when the platform, integration pattern, security constraints, document templates and reporting expectations are confirmed. Add schedule or trade scope when requirements introduce a public candidate portal, complex legacy migration, multiple offer-letter variants, e-signature, Microsoft Teams integration, advanced analytics, high-availability requirements, or material policy changes.')

doc.add_heading('Recommended Working Cadence',1)
cadence=[('Weekly','Backlog refinement and stakeholder decision session'),('Every 2 weeks','Demonstration, acceptance review and next-sprint commitment'),('Continuous','Automated unit/integration checks, audit logging and deployment pipeline'),('At each milestone','Formal exit review against the gate criteria above'),('Weeks 27-29','Structured UAT cycles with daily defect triage')]
table(doc,['Frequency','Activity'],cadence,[1.3,5.7],9)

doc.add_heading('Source References',1)
for s in ['Online Job Offer System - User Stories V2.xlsx','Job Offer Business Logic Rules.docx','Job Offer Programmer Detailed Flowchart - Editable Text.docx']:
    p=doc.add_paragraph(); p.style=styles['Normal']; r=p.add_run(s); r.bold=True
doc.add_paragraph('Prepared as a proposed delivery baseline. Dates should be converted from relative weeks to calendar dates after kickoff is approved.')

doc.core_properties.title='Online Job Offer System - Proposed Implementation Timeline'
doc.core_properties.subject='30-week delivery plan for one full-stack developer'
doc.core_properties.author='OpenAI Codex'
doc.save(OUT)
print(OUT)
